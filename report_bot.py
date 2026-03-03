import streamlit as st
from openai import OpenAI
from docx import Document
from fpdf import FPDF
import io

# --- 1. UI SETUP & THEME ---
def setup_ui():
    st.set_page_config(page_title="Academic Report Bot", page_icon="🎓", layout="wide")
    
    # Hide default menu and add custom styles
    st.markdown("""
        <style>
        #MainMenu {visibility: hidden;}
        footer {visibility: hidden;}
        .stApp {background-color: #f5f7f9;}
        
        /* Chat bubbles */
        .stChatMessage[data-testid="stChatMessage"]:nth-child(odd) {
            background-color: #e8f4f8; border-left: 5px solid #2980b9;
        }
        .stChatMessage[data-testid="stChatMessage"]:nth-child(even) {
            background-color: #ffffff; border-left: 5px solid #27ae60;
            box-shadow: 0 2px 5px rgba(0,0,0,0.05);
        }
        </style>
    """, unsafe_allow_html=True)
    
    with st.sidebar:
        st.header("🎓 TechReport AI")
        st.markdown("---")

# Call UI Setup immediately
setup_ui()

# --- 2. API KEY HANDLING ---
# Checks for key in Secrets first, otherwise asks user (Safe fallback)
if "OPENAI_API_KEY" in st.secrets:
    api_key = st.secrets["OPENAI_API_KEY"]
else:
    api_key = st.sidebar.text_input("Enter OpenAI API Key", type="password")
    if not api_key:
        st.warning("⚠️ Please enter an OpenAI API Key to proceed.")

# --- 3. SYSTEM PROMPTS ---
def get_system_prompt(level, field):
    if level == "Undergraduate":
        return f"You are a tutor for an Undergraduate student in {field}. Write a report focusing on fundamental concepts, definitions, and basic analysis. Structure: Intro, Core Concepts, Conclusion."
    elif level == "Masters":
        return f"You are a consultant for a Masters student in {field}. Write a report focusing on industry application, critical analysis, and methodology. Tone: Professional and analytical."
    elif level == "Ph.D.":
        return f"You are a research associate for a Ph.D. candidate in {field}. Write a highly rigorous report focusing on novelty, theoretical contribution, and literature context. Use academic jargon appropriate for the field."
    return f"You are a technical writer in {field}."

# --- 4. EXPORT FUNCTIONS ---
def create_docx(text):
    doc = Document()
    doc.add_heading('Technical Report', 0)
    for line in text.split('\n'):
        if line.strip():
            doc.add_paragraph(line)
    bio = io.BytesIO()
    doc.save(bio)
    return bio

def create_pdf(text):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.set_font("Arial", style="B", size=16)
    pdf.cell(0, 10, txt="Technical Report", ln=True, align='C')
    pdf.ln(10)
    pdf.set_font("Arial", size=12)
    
    # Sanitize text for PDF (Basic Latin-1 support)
    safe_text = text.encode('latin-1', 'ignore').decode('latin-1')
    pdf.multi_cell(0, 10, safe_text)
    return pdf.output(dest="S").encode("latin-1")

# --- 5. MAIN APP LOGIC ---
st.title("Academic Technical Report Generator")

# Sidebar Inputs
level = st.sidebar.selectbox("Qualification Level", ["Undergraduate", "Masters", "Ph.D."])
field = st.sidebar.text_input("Field of Study", "Computer Science")

# Chat History
if "messages" not in st.session_state:
    st.session_state.messages = []

for msg in st.session_state.messages:
    with st.chat_message(msg["role"]):
        st.markdown(msg["content"])

# Chat Input
if prompt := st.chat_input("Enter report topic..."):
    if not api_key:
        st.error("Stop: You need an API Key first.")
    else:
        # User Message
        st.session_state.messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.markdown(prompt)

        # AI Generation
        client = OpenAI(api_key=api_key)
        system_instruction = get_system_prompt(level, field)
        messages = [{"role": "system", "content": system_instruction}] + st.session_state.messages

        with st.chat_message("assistant"):
            msg_placeholder = st.empty()
            full_response = ""
            try:
                stream = client.chat.completions.create(
                    model="gpt-4",
                    messages=messages,
                    stream=True
                )
                for chunk in stream:
                    if chunk.choices[0].delta.content:
                        full_response += chunk.choices[0].delta.content
                        msg_placeholder.markdown(full_response + "▌")
                msg_placeholder.markdown(full_response)
                
                # Save and Show Download Options
                st.session_state.messages.append({"role": "assistant", "content": full_response})
                
                st.markdown("---")
                col1, col2 = st.columns(2)
                with col1:
                    st.download_button("📄 Download Word", create_docx(full_response).getvalue(), "report.docx")
                with col2:
                    st.download_button("📕 Download PDF", create_pdf(full_response), "report.pdf")

            except Exception as e:
                st.error(f"Error: {e}")
